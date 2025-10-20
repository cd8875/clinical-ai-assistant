import PyPDF2
import pdfplumber
from docx import Document
from typing import Dict
import re

class DocumentParser:
    
    @staticmethod
    def parse_pdf(file_path: str) -> Dict[str, str]:
        try:
            with pdfplumber.open(file_path) as pdf:
                text = ""
                for page in pdf.pages:
                    text += page.extract_text() or ""
                
                if len(text.strip()) > 100:
                    return {
                        "text": DocumentParser._clean_text(text),
                        "method": "pdfplumber",
                        "pages": len(pdf.pages)
                    }
        except Exception as e:
            print(f"pdfplumber failed: {e}")
        
        try:
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                text = ""
                for page in reader.pages:
                    text += page.extract_text() or ""
                
                return {
                    "text": DocumentParser._clean_text(text),
                    "method": "pypdf2",
                    "pages": len(reader.pages)
                }
        except Exception as e:
            raise ValueError(f"Failed to parse PDF: {e}")
    
    @staticmethod
    def parse_docx(file_path: str) -> Dict[str, str]:
        try:
            doc = Document(file_path)
            text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
            
            return {
                "text": DocumentParser._clean_text(text),
                "method": "python-docx",
                "paragraphs": len(doc.paragraphs)
            }
        except Exception as e:
            raise ValueError(f"Failed to parse DOCX: {e}")
    
    @staticmethod
    def parse_txt(file_path: str) -> Dict[str, str]:
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            
            return {
                "text": DocumentParser._clean_text(text),
                "method": "plain_text"
            }
        except Exception as e:
            raise ValueError(f"Failed to read text file: {e}")
    
    @staticmethod
    def _clean_text(text: str) -> str:
        text = re.sub(r'\s+', ' ', text)
        text = re.sub(r'[^\w\s\.\,\-\:\;\(\)\%\/\+\<\>\=]', '', text)
        text = text.replace('\n\n\n', '\n\n')
        return text.strip()
    
    @staticmethod
    def extract_metadata(text: str) -> Dict:
        metadata = {}
        
        patterns = {
            'patient_name': r'Patient Name[:\s]+([A-Za-z\s]+)',
            'patient_id': r'(?:Patient ID|MRN|Medical Record)[:\s]+([A-Z0-9\-]+)',
            'date': r'Date[:\s]+(\d{1,2}[\/\-]\d{1,2}[\/\-]\d{2,4})',
            'doctor': r'(?:Dr\.|Doctor|Physician)[:\s]+([A-Za-z\s\.]+)',
        }
        
        for key, pattern in patterns.items():
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                metadata[key] = match.group(1).strip()
        
        return metadata
    
    @staticmethod
    def parse_document(file_path: str, file_type: str) -> Dict:
        parsers = {
            'pdf': DocumentParser.parse_pdf,
            'docx': DocumentParser.parse_docx,
            'txt': DocumentParser.parse_txt
        }
        
        parser = parsers.get(file_type.lower())
        if not parser:
            raise ValueError(f"Unsupported file type: {file_type}")
        
        result = parser(file_path)
        result['metadata'] = DocumentParser.extract_metadata(result['text'])
        return result